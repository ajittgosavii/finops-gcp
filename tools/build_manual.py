"""Generate the technical manual (.docx) from this codebase.

    DATA_SOURCE=demo PYTHONPATH=services/api python tools/build_manual.py

Every inventory in this document is READ FROM THE CODE at build time: the 18
connectors and the secrets each needs, the 57 FOCUS columns, the 14 API routes,
the 11 typed agent tools and their docstrings, the 59 optimization levers, the
model ids. Nothing is typed by hand.

That is the point. A manual that lists endpoints by hand is wrong the first time
somebody adds one, and a wrong manual is worse than no manual -- the reader
trusts it. Here, adding a route adds a row; deleting a tool deletes one.

The diagrams are the same PNGs the deck embeds, rendered by tools/diagrams.py.
The .docx is a build artifact and is gitignored, like the deck.
"""

from __future__ import annotations

import os
import sys
from typing import List, Tuple

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, ROOT)
sys.path.insert(0, os.path.join(ROOT, "tools"))
sys.path.insert(0, os.path.join(ROOT, "packages", "finops-core", "src"))
sys.path.insert(0, os.path.join(ROOT, "services", "api"))
os.environ.setdefault("DATA_SOURCE", "demo")

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402

INK = RGBColor(0x0B, 0x14, 0x2A)
BODY = RGBColor(0x2E, 0x3A, 0x4E)
MUTED = RGBColor(0x4E, 0x57, 0x66)
ACCENT = RGBColor(0x1E, 0x6F, 0xD9)
CRIMSON = RGBColor(0xC2, 0x33, 0x33)

TITLE = "Multi-Cloud FinOps on Google Cloud"
SUBTITLE = "Technical Manual"
ORG = "Infosys · prepared for Con Edison"


# --------------------------------------------------------------------------
# Facts, read from the running code
# --------------------------------------------------------------------------


def facts():
    import build_deck as bd
    from finops_core import connectors, focus
    from finops_core.engines import optimize

    from app.agents.tools import build_tools
    from app.repository import DemoRepository
    from app.settings import Settings

    settings = Settings(data_source="demo")
    tools_by_agent = build_tools(DemoRepository(settings))
    tools = {}
    for agent, ts in tools_by_agent.items():
        for t in ts:
            tools.setdefault(t.__name__, {"doc": (t.__doc__ or "").strip().split("\n")[0], "agents": []})
            tools[t.__name__]["agents"].append(agent)

    from app.main import app as fastapi_app

    routes = sorted(
        {
            (sorted(r.methods - {"HEAD", "OPTIONS"})[0], r.path, (r.endpoint.__doc__ or "").strip().split("\n")[0])
            for r in fastapi_app.routes
            if hasattr(r, "methods") and r.path.startswith("/api")
        },
        key=lambda t: t[1],
    )

    return {
        "deck": bd.gather(),
        "specs": connectors.specs(),
        "schema": focus.SCHEMA,
        "mandatory": focus.MANDATORY_COLUMNS,
        "levers": optimize.LEVERS,
        "profiles": optimize._PROFILES,
        "tools": tools,
        "agents": list(tools_by_agent),
        "routes": routes,
        "settings": settings,
        "cost_per_question": bd.gemini_cost_per_question(),
    }


# --------------------------------------------------------------------------
# Layout primitives
# --------------------------------------------------------------------------


def _style_base(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BODY
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    for name, size, colour in (("Heading 1", 18, INK), ("Heading 2", 13, INK), ("Heading 3", 11, ACCENT)):
        st = doc.styles[name]
        st.font.name = "Segoe UI"
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True


def para(doc, text="", size=10, bold=False, colour=BODY, italic=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = colour
    p.paragraph_format.space_after = Pt(space_after)
    return p


def callout(doc, heading: str, text: str) -> None:
    """A bordered note. Used only where the reader would otherwise be misled."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    r = cell.paragraphs[0].add_run(heading)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = CRIMSON
    p = cell.add_paragraph()
    r2 = p.add_run(text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = BODY
    doc.add_paragraph()


def table(doc, headers: List[str], rows: List[Tuple], widths: List[float] = None, size=8.5) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(size)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v))
            r.font.size = Pt(size)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


def figure(doc, name: str, caption_text: str) -> None:
    import diagrams as dg

    png = os.path.join(dg.OUT_DIR, f"{name}.png")
    if not os.path.exists(png):
        dg.build_all()
    doc.add_picture(png, width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = para(doc, caption_text, size=8.5, italic=True, colour=MUTED)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, f"Editable vector source: docs/diagrams/{name}.svg", size=8, colour=MUTED)


def toc(doc) -> None:
    """A real TOC field. Word populates it on open (or F9); we cannot compute the
    page numbers here, so we ask Word to."""
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    run = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Right-click and choose “Update Field” to build the table of contents."
    run.append(t)
    fld.append(run)
    p._p.append(fld)
    doc.add_page_break()


# --------------------------------------------------------------------------
# The document
# --------------------------------------------------------------------------


def build(out: str) -> str:
    f = facts()
    d = f["deck"]
    doc = Document()
    _style_base(doc)

    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.9)

    # ---- cover -----------------------------------------------------------
    para(doc, ORG, size=10, colour=MUTED, space_after=40)
    para(doc, TITLE, size=26, bold=True, colour=INK, space_after=2)
    para(doc, SUBTITLE, size=15, colour=ACCENT, space_after=18)
    para(doc,
         "Four clouds — AWS, Azure, Google Cloud and OCI — normalised to the FinOps Foundation's "
         "FOCUS 1.2 specification, in one BigQuery warehouse, served by FastAPI on Cloud Run with a "
         "Google ADK agent team on Gemini.", size=11)
    para(doc, "", space_after=30)
    table(doc, ["", ""], [
        ("Clouds", "AWS · Azure · Google Cloud · OCI"),
        ("Connectors", f"{len(f['specs'])}"),
        ("FOCUS columns", f"{len(f['schema'])} ({len(f['mandatory'])} mandatory)"),
        ("Optimization levers", f"{len(f['levers'])}"),
        ("REST endpoints", f"{len(f['routes'])}"),
        ("Agent tools", f"{len(f['tools'])} typed, 0 SQL"),
        ("Reasoning model", f["settings"].model_reasoning),
        ("Routing model", f["settings"].model_routing),
    ], widths=[1.8, 4.6], size=9)
    doc.add_page_break()

    doc.add_heading("Contents", level=1)
    toc(doc)

    # ---- 1. Purpose ------------------------------------------------------
    doc.add_heading("1. Purpose and audience", level=1)
    para(doc,
         "This manual documents the platform as built. It is written for two readers: an architect who "
         "needs to review or extend the system, and an operator who needs to run it. Every inventory in "
         "it — connectors, columns, endpoints, tools, levers, model identifiers — is generated by reading "
         "the code at build time, so a section cannot silently fall out of date.")
    callout(doc, "What this document does not claim",
            "The figures throughout describe a synthetic 24-month utility estate shipped with the platform. "
            "They are not Con Edison's bill. Sign-in (IAP) is target state: it is not in the Terraform and the "
            "API ships today with no authentication. The OCI connector has been tested for shape but has never "
            "run against a live tenancy. AWS Cost Explorer exposes no list price, so on that ingest path "
            "ListCost is set equal to BilledCost and Effective Savings Rate is understated.")

    # ---- 2. System overview ---------------------------------------------
    doc.add_heading("2. System overview", level=1)
    para(doc,
         "Read the diagram downward. A vendor bill enters at the top and leaves at the bottom as an answer "
         "on someone's screen. Everything narrows to one table, and nothing above that table has ever seen "
         "a vendor-specific field.")
    figure(doc, "hld", "Figure 1 — High level design. Six layers: sources, identity, ingest, warehouse, serving, experience.")

    doc.add_heading("2.1 The architectural bet", level=2)
    para(doc,
         "Every source is normalised to FOCUS 1.2 on ingest. This buys three properties. First, a single "
         "definition of every number: Effective Savings Rate is computed once, in one function, from ListCost "
         "and EffectiveCost. Second, vendor neutrality by construction rather than by adapter — a procured "
         "FinOps tool becomes one Connector subclass and one registry line. Third, survivability: FOCUS leaves "
         "ProviderName a free string on purpose, so a provider the specification's authors never met still "
         "loads, allocates and forecasts.")

    # ---- 3. FOCUS --------------------------------------------------------
    doc.add_heading("3. The FOCUS contract", level=1)
    para(doc,
         f"finops_core.focus.SCHEMA defines {len(f['schema'])} columns, of which {len(f['mandatory'])} are "
         "mandatory. It is the only schema truth in the repository: the BigQuery DDL, the Terraform table "
         "definition and the ingest job all derive from it, and services/api/tests/test_schema_parity.py "
         "fails the build if any of them drift.")

    key = ["BilledCost", "EffectiveCost", "ListCost", "CommitmentDiscountStatus", "ProviderName",
           "SubAccountId", "ServiceCategory", "ChargeCategory", "ChargePeriodStart", "Tags"]
    by_name = {c.name: c for c in f["schema"]}
    table(doc, ["Column", "Level", "Type", "What it is for"],
          [(n, by_name[n].feature_level, by_name[n].dtype, _column_purpose(n)) for n in key if n in by_name],
          widths=[1.9, 0.9, 0.8, 3.2])

    callout(doc, "Why ProviderName is not an enum",
            "FOCUS deliberately leaves it a free string, so the specification survives a cloud its authors have "
            "not met. Code that consumes bill data must therefore tolerate a ProviderName it has never seen. "
            "optimize._profile() returns None for an unknown cloud and the rate detectors skip it, rather than "
            "raising KeyError deep inside a detector — which is exactly what happened when OCI was added.")

    # ---- 4. Sources and connectors ---------------------------------------
    doc.add_heading("4. Data sources and connectors", level=1)
    para(doc, f"{len(f['specs'])} connectors ship. Every SDK import is lazy, so importing a connector with "
              "nothing installed succeeds and test_connection() reports the missing package. A binding that "
              "fails contributes zero rows and a reason; the page still renders the clouds that are wired.")
    table(doc, ["Key", "Source", "Clouds", "FOCUS", "Required secrets"],
          [(s.key, s.display_name, ", ".join(s.clouds), s.focus_support,
            ", ".join(s.required_secrets) or "none") for s in f["specs"]],
          widths=[1.1, 1.7, 1.2, 0.7, 2.1], size=7.5)

    doc.add_heading("4.1 Connecting each cloud", level=2)
    para(doc, "One credential per payer, not one per account. A second credential means a second payer — which "
              "a regulated utility does have, because regulated and unregulated entities cannot share a bill.")
    figure(doc, "cloud_onboarding", "Figure 2 — What one credential covers, per cloud.")
    callout(doc, "OCI is the exception to the keyless story",
            "AWS and Azure are reached through Workload Identity Federation: nothing is stored and nothing needs "
            "rotating. There is no Workload Identity path from Google Cloud to OCI Object Storage — the OCI SDK "
            "signs each request with an RSA private key. That key lives in Secret Manager and is the one "
            "credential in this system that somebody must rotate. Further, OCI's cost reports live in a bucket "
            "Oracle owns, not yours: tenancy administrator is not sufficient, and you must add an IAM policy "
            "endorsing your group into Oracle's usage-report tenancy.")

    # ---- 5. Ingest -------------------------------------------------------
    doc.add_heading("5. Ingest", level=1)
    para(doc, "A Cloud Scheduler job triggers a Cloud Run Job nightly. It pulls each configured binding, "
              "normalises to FOCUS, validates, lands the raw Parquet in Cloud Storage, then loads to BigQuery "
              "through an idempotent staging table and a transactional partition replace.")
    para(doc, "Cloud Storage is not a cache. It is the replay source: a transform fixed six months from now is "
              "re-run against landed files rather than by re-pulling four vendors' bills.")

    # ---- 6. Warehouse ----------------------------------------------------
    doc.add_heading("6. Warehouse and cost guards", level=1)
    para(doc, "focus_costs is partitioned by DATE(ChargePeriodStart) and clustered by ProviderName, "
              "ServiceCategory and tag_application. opportunities holds a nightly snapshot of the row-level "
              "detectors; the API reads its MAX(as_of) partition.")
    table(doc, ["Guard", "Where it lives", "What happens when it bites"], [
        ("require_partition_filter = TRUE", "BigQuery DDL", "A query with no bound on the partition key FAILS rather than scanning two years."),
        ("maximum_bytes_billed", "every query job", "A runaway query FAILS rather than arriving as an invoice."),
        ("Nightly detector snapshot", "Cloud Run Job", "Row-level detectors never run per request; the answer changes once a day."),
    ], widths=[2.0, 1.4, 3.2], size=8.5)
    callout(doc, "The guards are declarative, not conventional",
            "They are properties of the table and of every job, not rules a developer must remember. Writing the "
            "partition guard immediately caught a query in our own repository that had no WHERE clause at all.")

    # ---- 7. API ----------------------------------------------------------
    doc.add_heading("7. REST API", level=1)
    para(doc, f"{len(f['routes'])} endpoints. Every scoped endpoint takes the same filter row (cloud, "
              "application, business unit, environment, period) and resolves it into a Scope value object.")
    table(doc, ["Method", "Path", "Purpose"],
          [(m, p, doc_ or "—") for m, p, doc_ in f["routes"]], widths=[0.7, 2.0, 3.9], size=8)
    callout(doc, "Dimensions are whitelisted, never interpolated",
            "A value in a query can be parameterised safely. A column name cannot — it must be interpolated into "
            "the SQL text. resolve_dimension() therefore checks every requested dimension against a GROUPABLE "
            "whitelist. An unchecked dimension string is an injection vector.")

    # ---- 8. Agents -------------------------------------------------------
    doc.add_heading("8. The agent layer", level=1)
    para(doc, f"A coordinator on {f['settings'].model_routing} routes; {len(f['agents'])} specialists "
              f"({', '.join(f['agents'])}) reason on {f['settings'].model_reasoning}. The specialists are held "
              "as AgentTool instances rather than sub_agents: transferring control means the last specialist to "
              "speak writes the final answer in its own register, and a FinOps answer must arrive in one voice, "
              "pitched at one persona.")

    doc.add_heading("8.1 Why the model never sees SQL", level=2)
    para(doc, "Google's ADK ships a BigQuery execute_sql toolset. It is deliberately not used. Given a SQL "
              "prompt, a language model will invent its own Effective Savings Rate — drop the on-demand "
              "denominator, count Purchase rows — and return a number that is plausible, wrong and uncaught. "
              f"Instead the model is given {len(f['tools'])} typed tools that call the same engine functions the "
              "REST endpoints call. It cannot compute. It can only ask.")
    table(doc, ["Tool", "Available to", "What it returns"],
          [(name, ", ".join(v["agents"]), v["doc"]) for name, v in sorted(f["tools"].items())],
          widths=[1.7, 1.5, 3.4], size=7.5)
    para(doc, f"Measured cost: ${f['cost_per_question']:.3f} per question "
              f"(~${f['cost_per_question'] * 4400:,.0f}/month at 4,400 questions).", size=9, colour=MUTED)

    doc.add_heading("8.2 One request, and one question", level=2)
    figure(doc, "lld", "Figure 3 — Both paths end at the same function, which is why the Copilot cannot contradict a dashboard.")
    figure(doc, "lld_technical", "Figure 4 — The same system at module level, for reviewers.")

    # ---- 9. Engines ------------------------------------------------------
    doc.add_heading("9. Analytics engines", level=1)
    para(doc, "All engines are pure pandas and numpy. They contain no Streamlit import, no cloud SDK and no "
              "framework coupling, which is what allowed roughly 9,100 lines to move from the Streamlit "
              "reference implementation into an installable package untouched.")
    table(doc, ["Module", "Responsibility", "Note"], [
        ("kpi", "Executive KPIs", "Every executive formula, defined exactly once."),
        ("forecast", "24-month forecast", "Method auto-selected by a WAPE backtest; overlays commitment-expiry cliffs."),
        ("budget", "Variance and run-rate", "Variance$ = Actual − Budget, exactly."),
        ("anomaly", "Anomaly detection", "STL decomposition plus a modified z-score on the residual, with a dollar floor."),
        ("allocation", "Showback and chargeback", "Names unallocated spend rather than silently spreading it."),
        ("optimize", f"{len(f['levers'])} levers", "Row-level detectors; runs nightly, never per request."),
    ], widths=[1.0, 1.7, 3.9], size=8.5)

    doc.add_heading("9.1 Per-cloud rate profiles", level=2)
    para(doc, "Rate levers need to know a cloud's commitment instrument and its interruptible-capacity discount. "
              "The lookup is total: a cloud without a profile receives no rate lever, because we will not invent "
              "a commitment instrument we cannot name. It still receives every provider-agnostic finding.")
    table(doc, ["Cloud", "Commitment rate", "Interruptible discount", "Commitment lever"],
          [(c, f"{p.commitment_rate:.0%}", f"{p.spot_discount:.0%}", p.commitment_lever)
           for c, p in f["profiles"].items()], widths=[1.2, 1.5, 1.7, 1.5], size=8.5)

    # ---- 10. End user ----------------------------------------------------
    doc.add_heading("10. End user view", level=1)
    figure(doc, "end_user_view", "Figure 5 — Five personas, and the pages that answer their question.")
    para(doc, "One scope governs every panel on a page, so two charts on the same screen cannot disagree. Every "
              "chart has a table twin, and every table downloads as CSV: no value is reachable only through a "
              "tooltip.")

    # ---- 11. Security ----------------------------------------------------
    doc.add_heading("11. Security model", level=1)
    table(doc, ["Concern", "Position"], [
        ("Cloud access", "Read-only, everywhere."),
        ("AWS / Azure", "Workload Identity Federation. No static key is stored."),
        ("Google Cloud", "Application Default Credentials on the service account."),
        ("OCI", "One RSA signing key in Secret Manager. It must be rotated."),
        ("Model credentials", "Vertex via ADC. No model API key exists anywhere."),
        ("SQL", "Dimensions whitelisted; values parameterised; the model never writes a query."),
        ("Cost", "Partition filter required and bytes-billed capped, in the DDL and on every job."),
        ("Authentication", "TARGET STATE — IAP is not in the Terraform; the API ships with no auth."),
    ], widths=[1.6, 5.0], size=8.5)

    # ---- 12. Deployment and operations -----------------------------------
    doc.add_heading("12. Deployment and operations", level=1)
    table(doc, ["Component", "Runtime", "Trigger"], [
        ("services/api", "Cloud Run (scales to zero)", "HTTPS"),
        ("services/ingest", "Cloud Run Job", "Cloud Scheduler, nightly 03:15 America/New_York"),
        ("opportunities snapshot", "Cloud Run Job", "Nightly, after ingest"),
        ("web", "Static build behind the load balancer", "—"),
        ("Infrastructure", "Terraform (infra/terraform)", "terraform apply"),
        ("CI", "GitHub Actions", "push / pull request"),
    ], widths=[1.7, 2.3, 2.6], size=8.5)

    doc.add_heading("12.1 Runbook", level=2)
    table(doc, ["Symptom", "Likely cause", "Action"], [
        ("A cloud's rows are missing", "Binding failed; it contributes zero rows and a reason.",
         "Check the Integrations page, then the Job logs. Other clouds keep rendering."),
        ("OCI returns nothing", "The endorse policy is missing, or the key expired.",
         "Confirm 'endorse group … to read objects in tenancy usage-report'. Rotate the key."),
        ("A query fails on bytes billed", "The guard bit. This is intended.",
         "Narrow the period. Do not raise the cap without understanding the scan."),
        ("Effective Savings Rate looks low", "Ingested via Cost Explorer, which has no list price.",
         "Switch that binding to a FOCUS Data Export. ESR is understated, not wrong."),
        ("The Copilot declines to answer", "No tool can supply the figure.",
         "By design. It says what data would be needed rather than inventing a number."),
    ], widths=[1.9, 2.1, 2.6], size=8)

    # ---- 13. Testing -----------------------------------------------------
    doc.add_heading("13. Testing", level=1)
    para(doc, "No test requires a GCP project, a credential or a network. The BigQuery paths are exercised "
              "against a fake client that asserts every query prunes the partition, caps bytes billed and "
              "parameterises its filters.")
    table(doc, ["Suite", "Command", "Covers"], [
        ("finops-core", "pytest packages/finops-core", "FOCUS schema, KPIs, engines, connectors, unknown providers"),
        ("api", "cd services/api && DATA_SOURCE=demo pytest", "Routes, scope, SQL whitelist, cost guards, schema parity"),
        ("ingest", "cd services/ingest && pytest", "Idempotent staging, transactional partition replace"),
        ("diagrams & deck", "DATA_SOURCE=demo PYTHONPATH=services/api pytest tools", "Counts on diagrams match code; nothing off-slide; every slide has notes"),
        ("frontend", "cd web && npx tsc --noEmit && npm run build", "Types and bundle"),
        ("infrastructure", "terraform -chdir=infra/terraform validate", "Configuration"),
    ], widths=[1.3, 2.5, 2.8], size=8)

    # ---- 14. Limits ------------------------------------------------------
    doc.add_heading("14. Known limits", level=1)
    for t in (
        "The estate behind every figure in this manual is synthetic. It is not Con Edison's bill.",
        "Authentication is target state. IAP is not in the Terraform and the API ships with no auth.",
        "The OCI connector has never run against a live tenancy. Its shape is tested; its network path is not.",
        "AWS Cost Explorer exposes no list price. On that path ListCost equals BilledCost and ESR is understated; "
        "use a FOCUS Data Export.",
        "A cloud bill cannot contain a budget or a business driver. budgets() and drivers() return empty frames on "
        "purpose rather than inventing plausible numbers.",
        "gcloud is not installed on the build machine, so nothing in this repository has been validated against a "
        "live GCP project.",
    ):
        p = doc.add_paragraph(t, style="List Bullet")
        p.runs[0].font.size = Pt(9.5)

    doc.save(out)
    return out


def _column_purpose(name: str) -> str:
    return {
        "BilledCost": "What appeared on the invoice this period.",
        "EffectiveCost": "Amortised cost. Every executive KPI uses this.",
        "ListCost": "On-demand-equivalent price. The ESR denominator.",
        "CommitmentDiscountStatus": "Used or Unused. 'Unused' is waste the bill states outright.",
        "ProviderName": "Free string, by design. Not an enum.",
        "SubAccountId": "The account beneath the payer: account, subscription, project or compartment.",
        "ServiceCategory": "Closed enum. Compute, Storage, Databases, and so on.",
        "ChargeCategory": "Closed enum: Usage, Purchase, Tax, Credit, Adjustment.",
        "ChargePeriodStart": "The BigQuery partition key.",
        "Tags": "Stored as a JSON string; six canonical tag_* columns are materialised on ingest.",
    }.get(name, "")


def main() -> None:
    out = os.path.join(ROOT, "docs", "Infosys_FinOps_GCP_Technical_Manual.docx")
    build(out)
    print("wrote", os.path.relpath(out, ROOT))


if __name__ == "__main__":
    main()
