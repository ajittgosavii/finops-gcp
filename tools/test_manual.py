"""The technical manual must be generated, not written.

A manual that lists endpoints by hand is wrong the first time somebody adds one,
and a wrong manual is worse than no manual because the reader trusts it. Every
inventory in this document is read from the code at build time; these tests
assert that the generated artifact actually contains what the code says.
"""

from __future__ import annotations

import os
import sys

import pytest

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, os.path.join(ROOT, "tools"))
sys.path.insert(0, os.path.join(ROOT, "packages", "finops-core", "src"))
sys.path.insert(0, os.path.join(ROOT, "services", "api"))
os.environ.setdefault("DATA_SOURCE", "demo")

pytest.importorskip("docx")
pytest.importorskip("matplotlib")

import build_manual as bm  # noqa: E402
from docx import Document  # noqa: E402


@pytest.fixture(scope="module")
def manual(tmp_path_factory):
    out = str(tmp_path_factory.mktemp("manual") / "m.docx")
    bm.build(out)
    return Document(out)


def _tables(doc):
    return [[[c.text.strip() for c in r.cells] for r in t.rows] for t in doc.tables]


def test_the_manual_embeds_every_diagram(manual) -> None:
    images = [r for r in manual.part.rels.values() if "image" in r.reltype]
    assert len(images) == 5, f"expected 5 figures, found {len(images)}"


def test_the_connector_table_lists_every_connector(manual) -> None:
    from finops_core import connectors

    rows = next(t for t in _tables(manual) if t[0][0] == "Key")
    assert len(rows) - 1 == len(connectors.REGISTRY)
    assert any(r[0] == "oci_native" for r in rows)


def test_the_api_table_lists_every_endpoint_with_a_description(manual) -> None:
    from app.main import app as fastapi_app

    endpoints = [r for r in fastapi_app.routes if hasattr(r, "methods") and r.path.startswith("/api")]
    rows = next(t for t in _tables(manual) if t[0][:2] == ["Method", "Path"])
    assert len(rows) - 1 == len(endpoints)
    assert not [r for r in rows[1:] if r[2] in ("", "—")], "an endpoint reached the manual undocumented"


def test_the_tool_table_lists_every_agent_tool(manual) -> None:
    from app.agents.tools import build_tools
    from app.repository import DemoRepository
    from app.settings import Settings

    tools = build_tools(DemoRepository(Settings(data_source="demo")))
    distinct = {t.__name__ for ts in tools.values() for t in ts}
    rows = next(t for t in _tables(manual) if t[0][0] == "Tool")
    assert len(rows) - 1 == len(distinct)


def test_the_profile_table_covers_every_profiled_cloud(manual) -> None:
    from finops_core.engines import optimize

    rows = next(t for t in _tables(manual) if t[0][0] == "Cloud")
    assert {r[0] for r in rows[1:]} == set(optimize._PROFILES)


def test_the_manual_states_its_limits(manual) -> None:
    """The section a reader most needs, and the one most likely to be quietly
    dropped. It must name all four."""
    text = "\n".join(p.text for p in manual.paragraphs) + "\n".join(
        c.text for t in manual.tables for r in t.rows for c in r.cells
    )
    for claim in ("synthetic", "IAP", "live tenancy", "Cost Explorer"):
        assert claim in text, f"the manual no longer discloses: {claim}"
